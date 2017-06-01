#!/usr/bin/env python
# -*-coding:utf-8-*-

from django.shortcuts import render
from django.http import HttpResponse
from django.conf import settings
from bson.objectid import ObjectId

from docx import Document
from docx.shared import Inches

import zujuan.dbconn as dbconn
import time

""" 整个系统的所有接口的功能都写在这里了
    包含有1、页面上所见的所有可点击的链接
    后面实现的功能，具体的见下面的描述
"""


def query_data(tabname, key=None):
    # 通用数据查询函数，返回列表形式的数据，key 指的是字段名
    # 如果不指定，就返回所有记录，如果指定了，仅返回 key 字段的值
    client, db, coll = dbconn.db(tabname)
    data = coll.find()
    if key:
        datas = [x.get(key) for x in data]
    else:
        datas = data
    client.close()
    return datas


def query_qtype_data():
    # 查询题目类型数据
    return query_data('qtype', 'name')


def query_qdiff_data():
    # 查询题目难易度数据
    return query_data('qdiff', 'name')


def query_qknow_data():
    # 查询知识点数据
    return query_data('qknow', 'zsd')


def query_question_data():
    # 查询问题数据
    return query_data('question')


def delete_data(tabname, qid):
    # 删除数据
    client, db, coll = dbconn.db(tabname)
    coll.delete_one({'_id': ObjectId(qid)})


def index(request):
    # 首页，列出来一些基本的功能属性
    return render(request, 'zujuan/index.html')


def question_type(request):
    # 问题类型管理接口
    # 主要做的是对问题类型的读取和添加
    names = query_qtype_data()

    # 创建题目类型，包含有 单选、多选、填空、判断、问答
    if request.method == 'GET':  # 渲染出来已有的问题类型列表
        return render(request, 'zujuan/question_type.html', {'names': names})
    else:  # POST 方法，生成/修改一个类型
        post_data = request.POST
        type_name = post_data.get('type_name')  # 试卷类型名
        if type_name in names:
            return render(request, 'zujuan/question_type.html', {'names': names})

        client, db, coll = dbconn.db('qtype')
        coll.insert_one({'name': type_name})
        names.append(type_name)
        return render(request, 'zujuan/question_type.html', {'names': names})


def question_difficult(request):
    # 创建题目的难易度列表
    names = query_qdiff_data()

    if request.method == 'GET':  # 渲染出来已有的问题难易度列表
        return render(request, 'zujuan/question_difficult.html', {'names': names})
    else:  # POST 方法，生成/修改一个类型
        post_data = request.POST
        type_name = post_data.get('type_name', '').strip()
        if type_name in names:
            return render(request, 'zujuan/question_difficult.html', {'names': names})

        client, db, coll = dbconn.db('qdiff')
        coll.insert_one({'name': type_name})
        names.append(type_name)
        return render(request, 'zujuan/question_difficult.html', {'names': names})


def question_knowledge(request):
    # 知识点内容的读取
    tabname = 'qknow'
    client, db, coll = dbconn.db(tabname)
    data = coll.find()
    zsd = [x.get('zsd') for x in data]  # 知识点

    if request.method == 'GET':  # 渲染出来已有的问题类型列表
        return render(request, 'zujuan/question_kg.html', {'datas': zsd})


def question_knowledge_post(request):
    zsd = query_qknow_data()

    # 知识点创建与维护
    if request.method == 'GET':  # 渲染出来已有的问题类型列表
        return render(request, 'zujuan/question_kg_post.html')
    else:  # POST 方法，生成/修改一个知识点
        post_data = request.POST
        content = post_data.get('content', '').strip()  # 知识点

        client, db, coll = dbconn.db('qknow')
        coll.insert_one({'zsd': content})
        zsd.append(content)
        return render(request, 'zujuan/question_kg.html', {'datas': zsd})


def question_list(request):
    # 题目列表读取接口，返回目前已创建的题目
    questions = query_question_data()
    questions = list(questions)
    for idx, val in enumerate(questions[:]):
        qid = str(val.get('_id'))
        val['id'] = qid
        questions[idx] = val

    return render(request, 'zujuan/question_list.html', {'questions': questions})


def question_create(request):
    # 创建一个题目，创建一个题目依赖于问题类型、难易度、知识点等数据
    # 故此处需要读取那些数据出来
    qtype = query_qtype_data()
    qdiff = query_qdiff_data()
    zds = query_qknow_data()

    if request.method == 'GET':  # 渲染出来已有的问题类型列表
        return render(request, 'zujuan/question_create.html', {'qtypes': qtype, 'qdiffs': qdiff, 'zsds': zds})
    else:  # POST 方法，生成/修改一个类型
        post_data = request.POST
        qtype = post_data.get('qtype')  # 问题类型
        qdiff = post_data.get('qdiff')  # 问题难度
        qscore = post_data.get('qscore')  # 问题分数
        qkg = post_data.getlist('qkg')  # 问题知识点
        qdesc = post_data.get('qdesc', '').strip()  # 问题题干
        qanswer = post_data.get('qanswer', '').strip()  # 问题答案

        client, db, coll = dbconn.db('question')

        ins_data = {
            'qtype': qtype,
            'qdiff': qdiff,
            'qscore': qscore,
            'qknow': qkg,
            'qdesc': qdesc,
            'qanswer': qanswer,
        }

        coll.insert_one(ins_data)
        return render(request, 'zujuan/question_ok.html')


def question_detail(request, qid):
    # 题目详情信息
    tabname = 'question'
    client, db, coll = dbconn.db(tabname)
    data = coll.find_one({'_id': ObjectId(qid)})
    data['qid'] = qid

    if request.method == 'GET':  # 根据 id 读取问题数据
        return render(request, 'zujuan/question_detail.html', {'q': data})


def question_edit(request, qid):
    # 题目的编辑
    tabname = 'question'
    client, db, coll = dbconn.db(tabname)
    data = coll.find_one({'_id': ObjectId(qid)})
    data['qid'] = qid

    qtype = query_qtype_data()
    qdiff = query_qdiff_data()
    zsd = query_qknow_data()

    qknows = data.get('qknow')

    selected_qknows = []
    unselected = []
    for qk in zsd:
        if qk in qknows:
            selected_qknows.append(qk)
        else:
            unselected.append(qk)

    if request.method == 'GET':  # 根据 id 读取问题数据
        return render(request, 'zujuan/question_edit.html',
                      {'q': data, 'qtypes': qtype, 'qdiffs': qdiff, 'selected': selected_qknows,
                       'unselect': unselected})
    else:
        post_data = request.POST
        qtype = post_data.get('qtype')  # 问题类型
        qdiff = post_data.get('qdiff')  # 问题难度
        qscore = post_data.get('qscore')  # 问题分数
        qkg = post_data.getlist('qkg')  # 问题知识点
        qdesc = post_data.get('qdesc', '').strip()  # 问题题干
        qanswer = post_data.get('qanswer', '').strip()  # 问题答案

        client, db, coll = dbconn.db('question')

        ins_data = {
            'qtype': qtype,
            'qdiff': qdiff,
            'qscore': qscore,
            'qknow': qkg,
            'qdesc': qdesc,
            'qanswer': qanswer,
        }

        coll.update_one({'_id': ObjectId(qid)}, {'$set': ins_data})
        return render(request, 'zujuan/question_ok.html')


def question_delete(request, qid):
    # 题目删除
    tabname = 'question'
    delete_data(tabname, qid)
    return render(request, 'zujuan/delete_ok.html')


def qtype_delete(request, qid):
    # 题目类型的删除，暂未使用
    tabname = 'qtype'
    delete_data(tabname, qid)
    return render(request, 'zujuan/delete_ok.html')


def auto_zujuan(request):
    # 自动组卷
    # 组卷就是根据用户提交的知识点、难易度，找出来相关的题目数据
    # 将他们按题目类型分类后，再返回显示到页面上，同时将数据写入 docx 文件中
    tabname = 'question'

    questions = list(query_question_data())

    qknows = [x.get('qknow') for x in questions]
    ret_qk = []
    for qk in qknows:
        ret_qk.extend(qk)
    ret_qk = set(ret_qk)

    qdiffs = [x.get('qdiff') for x in questions]

    if request.method == 'GET':
        return render(request, 'zujuan/auto_zujuan.html', {'qknows': ret_qk, 'qdiffs': qdiffs})
    else:
        post_data = request.POST
        qdiff = post_data.get('qdiff')  # 难度指示
        qkgs = post_data.getlist('qkg')  # 知识点列表

        # 组卷的过程就是根据知识点和难度，从数据库中查询数据，生成一个题目列表
        client, db, coll = dbconn.db(tabname)

        data = []
        q1 = {}
        if qdiff != '全部':
            q1['qdiff'] = qdiff

        for qkg in qkgs:
            query = {}
            query.update(q1)
            query['qknow'] = {'$regex': qkg}
            tmp = coll.find(query)
            tmp = list(tmp)
            data.extend(tmp)

        # 过滤掉重复查询的题目数据
        ret_data = []
        qids = []
        for idx, val in enumerate(data):
            qid = str(val.get('_id'))
            if qid in qids:
                continue
            else:
                qids.append(qid)
                ret_data.append(val)

        ret_data = format_zujuan_result(ret_data)

        file = generate_doc(ret_data)

        return render(request, 'zujuan/zujuan_result.html', {'questions': ret_data, 'file': file})


def format_zujuan_result(datas):
    # 将查询出来的组卷结果，按照题目类型进行分组
    tmp_type = []
    ret_data = {}
    for data in datas:
        dtype = data.get('qtype')
        if dtype in tmp_type:
            ret_data[dtype].append(data)
        else:
            tmp_type.append(dtype)
            ret_data[dtype] = [data]
    return ret_data


def generate_doc(datas):
    # 生成 doc 文档

    document = Document()
    document.add_heading('试卷结果', 0)

    p = document.add_paragraph('本试卷为系统自动生成')

    for qtype, data in datas.items():
        document.add_heading(qtype, level=1)
        for d in data:
            document.add_heading(d.get('qdesc') + '(分数:' + str(d.get('qscore')) + ')', level=2)
            document.add_paragraph('难度:' + d.get('qdiff'))
            document.add_heading('知识点', level=3)
            qkg = d.get('qknow')
            if isinstance(qkg, type([])):
                for kg in qkg:
                    document.add_paragraph(kg, style='ListNumber')
            else:
                document.add_paragraph(qkg, style='ListNumber')

    path = settings.MEDIA_ROOT
    file_name = str(int(time.time())) + '.docx'
    file = path + file_name
    document.save(file)
    return file_name
